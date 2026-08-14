"""Extract raw text from legal documents (.pdf, .docx, .txt), with OCR fallback."""

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from backend.config import SCANNED_PDF_THRESHOLD
from backend.ocr import TESSERACT_MISSING_MESSAGE, ocr_available, ocr_pdf

SUPPORTED_EXTENSIONS: tuple[str, ...] = (".pdf", ".docx", ".txt")

_MIN_USABLE_CHARS = 10
_NO_TEXT_MESSAGE = (
    "No readable text found — the document may be scanned (image-only) or empty."
)


@dataclass(frozen=True)
class Extraction:
    """The text of a document plus how it was obtained."""

    text: str
    used_ocr: bool = False
    pages: int | None = None


def extract_document(file_path: str) -> Extraction:
    """Extract text from a .pdf, .docx, or .txt file, using OCR for scanned PDFs."""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise ValueError(f"Not a file: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_document(path)
    if suffix == ".docx":
        return Extraction(_finalize(_extract_docx(path)))
    if suffix == ".txt":
        return Extraction(_finalize(_extract_txt(path)))
    if suffix == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported. Open the file in Word and save "
            "it as .docx, then try again."
        )
    raise ValueError(
        f"Unsupported file type: {suffix or '(none)'}. "
        f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


def extract_text(file_path: str) -> str:
    """Extract raw text from a .pdf, .docx, or .txt file."""
    return extract_document(file_path).text


def _extract_pdf_document(path: Path) -> Extraction:
    """Read a PDF's text layer, falling back to OCR when the PDF is a scan."""
    embedded, page_count = _extract_pdf(path)
    embedded = _normalize(embedded)

    if len(embedded) >= SCANNED_PDF_THRESHOLD:
        return Extraction(_finalize(embedded, already_normalized=True), False, page_count)

    # Little or no text layer: this is almost certainly a scan.
    if not ocr_available():
        raise ValueError(TESSERACT_MISSING_MESSAGE)

    scanned = _normalize(ocr_pdf(str(path)))
    if len(scanned) < _MIN_USABLE_CHARS:
        raise ValueError(_NO_TEXT_MESSAGE)

    # If the text layer held a little text, keep whichever source gave more.
    if len(embedded) > len(scanned):
        return Extraction(embedded, False, page_count)
    return Extraction(scanned, True, page_count)


def _extract_pdf(path: Path) -> tuple[str, int]:
    """Return the PDF's embedded text and its page count. Empty text means a scan."""
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ValueError(f"Could not open the PDF — it may be corrupt. ({exc})") from exc

    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                raise ValueError(
                    "This PDF is password-protected. Remove the password and try again."
                )
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(
                "This PDF is encrypted and could not be opened. Remove the password "
                "and try again."
            ) from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            content = page.extract_text() or ""
        except Exception:
            content = ""  # A single unreadable page should not sink the document.
        if content.strip():
            pages.append(content.strip())

    return "\n\n".join(pages), len(reader.pages)


def _extract_docx(path: Path) -> str:
    """Extract text from all paragraphs and table cells of a DOCX."""
    try:
        document = Document(str(path))
    except Exception as exc:
        raise ValueError(
            f"Could not open the Word document — it may be corrupt or not a real "
            f".docx file. ({exc})"
        ) from exc

    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    if not blocks:
        raise ValueError(_NO_TEXT_MESSAGE)
    return "\n\n".join(blocks)


def _extract_txt(path: Path) -> str:
    """Read a plain-text file as UTF-8, falling back to latin-1."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="latin-1")
        except Exception as exc:
            raise ValueError(f"Could not read the text file. ({exc})") from exc
    except Exception as exc:
        raise ValueError(f"Could not read the text file. ({exc})") from exc


def _finalize(text: str, already_normalized: bool = False) -> str:
    """Normalize whitespace and reject documents with too little readable text."""
    result = text if already_normalized else _normalize(text)
    if len(result) < _MIN_USABLE_CHARS:
        raise ValueError(_NO_TEXT_MESSAGE)
    return result


def _normalize(text: str) -> str:
    """Collapse runs of blank lines and trailing spaces, and strip the result."""
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
